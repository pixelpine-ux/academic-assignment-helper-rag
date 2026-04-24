"""Script to create test fixtures - run inside Docker container"""
from docx import Document
from PyPDF2 import PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

# Create DOCX
doc = Document()
doc.add_paragraph("This is a sample DOCX document for testing.")
doc.add_paragraph("It contains multiple paragraphs.")
doc.add_paragraph("The parser should extract all text content properly.")
doc.save("sample.docx")

# Create PDF
packet = io.BytesIO()
can = canvas.Canvas(packet, pagesize=letter)
can.drawString(100, 750, "This is a sample PDF document for testing.")
can.drawString(100, 730, "It contains text content.")
can.drawString(100, 710, "The parser should extract all text properly.")
can.save()
packet.seek(0)

with open("sample.pdf", "wb") as f:
    f.write(packet.read())

print("Fixtures created successfully")
